
# data_converter_app.py
import streamlit as st
import pandas as pd
import numpy as np
from io import BytesIO
import plotly.express as px
import base64
from datetime import datetime
import time
import json
import seaborn as sns
import matplotlib.pyplot as plt
import docx
from docx import Document
import PyPDF2
import xml.etree.ElementTree as ET
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table
import mimetypes
import chardet
import pickle
import feather
import h5py
import yaml
import toml
import sqlite3
import csv
import zipfile
import tarfile
import gzip
import bz2
import lzma

# Page configuration
st.set_page_config(
    page_title="DataForge Omni",
    page_icon="🌌",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Enhanced CSS with animations
st.markdown("""
    <style>
    .app-wrapper {
        background: linear-gradient(135deg, #0a0e21 0%, #1e2a44 100%);
        min-height: 100vh;
        padding: 40px;
        color: #e2e8f0;
        animation: fadeIn 1s ease-in;
    }
    .header {
        background: linear-gradient(90deg, #ec4899, #8b5cf6);
        padding: 30px;
        border-radius: 15px;
        text-align: center;
        margin-bottom: 40px;
        animation: glowPulse 2.5s infinite alternate;
    }
    .header-title {
        color: #fff;
        font-size: 52px;
        font-weight: 900;
        margin: 0;
        text-shadow: 0 4px 8px rgba(0,0,0,0.3);
        animation: zoomIn 1s ease;
    }
    .header-subtitle {
        color: #ddd6fe;
        font-size: 22px;
        margin: 10px 0 0;
        animation: fadeInUp 1s ease 0.5s;
        animation-fill-mode: backwards;
    }
    .panel {
        background: #1f2a44;
        border-radius: 12px;
        padding: 25px;
        margin-bottom: 25px;
        border: 2px solid #4b5563;
        transition: all 0.3s ease;
        animation: slideUp 0.7s ease;
    }
    .panel:hover {
        border-color: #a78bfa;
        box-shadow: 0 8px 25px rgba(167, 139, 250, 0.3);
        transform: translateY(-3px);
    }
    .stButton>button {
        background: #4b5563;
        color: #e2e8f0;
        border-radius: 8px;
        border: none;
        padding: 12px 24px;
        transition: all 0.3s ease;
        font-weight: 700;
    }
    .stButton>button:hover {
        background: #a78bfa;
        color: #fff;
        transform: scale(1.05);
        box-shadow: 0 5px 15px rgba(167, 139, 250, 0.4);
    }
    .sidebar .sidebar-content {
        background: #1f2a44;
        color: #e2e8f0;
        animation: slideInLeft 0.6s ease;
    }
    .glow-text {
        animation: glowText 2s infinite alternate;
    }
    .action-panel {
        animation: pulse 1.8s infinite;
    }
    @keyframes fadeIn {
        from {opacity: 0;}
        to {opacity: 1;}
    }
    @keyframes glowPulse {
        from {box-shadow: 0 0 10px rgba(236, 72, 153, 0.5);}
        to {box-shadow: 0 0 20px rgba(167, 139, 250, 0.7);}
    }
    @keyframes zoomIn {
        from {transform: scale(0.8); opacity: 0;}
        to {transform: scale(1); opacity: 1;}
    }
    @keyframes fadeInUp {
        from {transform: translateY(20px); opacity: 0;}
        to {transform: translateY(0); opacity: 1;}
    }
    @keyframes slideUp {
        from {transform: translateY(50px); opacity: 0;}
        to {transform: translateY(0); opacity: 1;}
    }
    @keyframes slideInLeft {
        from {transform: translateX(-100%); opacity: 0;}
        to {transform: translateX(0); opacity: 1;}
    }
    @keyframes glowText {
        from {text-shadow: 0 0 5px #a78bfa;}
        to {text-shadow: 0 0 15px #ec4899;}
    }
    @keyframes pulse {
        0% {transform: scale(1);}
        50% {transform: scale(1.02);}
        100% {transform: scale(1);}
    }
    .spinner {
        animation: spin 1.2s linear infinite;
        display: inline-block;
        font-size: 24px;
    }
    @keyframes spin {
        0% {transform: rotate(0deg);}
        100% {transform: rotate(360deg);}
    }
    </style>
""", unsafe_allow_html=True)

# Initialize session state
for key in ['raw_df', 'processed_df', 'history', 'undo_stack', 'presets', 'file_metadata']:
    if key not in st.session_state:
        st.session_state[key] = [] if key in ['undo_stack', 'history'] else None
if st.session_state.presets is None:
    st.session_state.presets = {}
if st.session_state.file_metadata is None:
    st.session_state.file_metadata = {}

# Main UI
st.markdown('<div class="app-wrapper">', unsafe_allow_html=True)
st.markdown("""
    <div class="header">
        <h1 class="header-title">🌌 DataForge Omni</h1>
        <p class="header-subtitle">Universal data processing for every format imaginable</p>
    </div>
""", unsafe_allow_html=True)

# Sidebar for upload and options
with st.sidebar:
    st.markdown('<h2 class="glow-text">Data Input</h2>', unsafe_allow_html=True)
    upload_mode = st.radio("Upload Mode", ["Single", "Batch"], horizontal=True)
    files = st.file_uploader(
        "Upload Any File Type",
        type=None,  # Accept all files
        accept_multiple_files=upload_mode == "Batch",
        key="omni_upload"
    )
    fallback_mode = st.selectbox("Fallback for Unknown Files", ["Text", "Binary", "Skip"], index=0)

def load_data(file):
    """Load any file into a DataFrame with comprehensive format support"""
    with st.spinner("Processing <span class='spinner'>🌍</span>"):
        time.sleep(0.5)
        try:
            ext = file.name.split('.')[-1].lower() if '.' in file.name else ''
            mime, _ = mimetypes.guess_type(file.name)
            raw_data = file.read()
            
            # Store metadata
            st.session_state.file_metadata[file.name] = {
                "size": len(raw_data),
                "mime": mime or "application/octet-stream",
                "ext": ext,
                "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            }

            # Structured formats
            if ext == 'csv':
                return pd.read_csv(BytesIO(raw_data), encoding='utf-8', on_bad_lines='skip')
            elif ext in ['xlsx', 'xls']:
                return pd.read_excel(BytesIO(raw_data))
            elif ext == 'json':
                return pd.read_json(BytesIO(raw_data))
            elif ext == 'parquet':
                return pd.read_parquet(BytesIO(raw_data))
            elif ext == 'feather':
                return feather.read_dataframe(BytesIO(raw_data))
            elif ext == 'pickle' or ext == 'pkl':
                return pd.read_pickle(BytesIO(raw_data))
            elif ext == 'h5' or ext == 'hdf5':
                with h5py.File(BytesIO(raw_data), 'r') as f:
                    keys = list(f.keys())
                    return pd.read_hdf(BytesIO(raw_data), key=keys[0] if keys else 'data')
            elif ext == 'sql':
                conn = sqlite3.connect(':memory:')
                conn.cursor().executescript(raw_data.decode('utf-8'))
                return pd.read_sql_query("SELECT * FROM sqlite_master", conn)
            elif ext == 'xml':
                tree = ET.ElementTree(ET.fromstring(raw_data))
                root = tree.getroot()
                data = [{child.tag: child.text for child in elem} for elem in root]
                return pd.DataFrame(data)
            elif ext == 'yaml' or ext == 'yml':
                data = yaml.safe_load(raw_data)
                return pd.json_normalize(data)
            elif ext == 'toml':
                data = toml.loads(raw_data.decode('utf-8'))
                return pd.json_normalize(data)

            # Text-based formats
            elif ext == 'txt':
                encoding = chardet.detect(raw_data)['encoding'] or 'utf-8'
                return pd.DataFrame({"content": raw_data.decode(encoding).splitlines()})
            elif ext == 'md':
                encoding = chardet.detect(raw_data)['encoding'] or 'utf-8'
                return pd.DataFrame({"markdown": raw_data.decode(encoding).splitlines()})
            elif ext == 'html':
                return pd.read_html(BytesIO(raw_data))[0] if pd.read_html(BytesIO(raw_data)) else pd.DataFrame({"html": [raw_data.decode('utf-8')]})
            elif ext == 'docx':
                doc = Document(BytesIO(raw_data))
                text = [p.text for p in doc.paragraphs if p.text]
                return pd.DataFrame({"content": text})
            elif ext == 'pdf':
                pdf_reader = PyPDF2.PdfReader(BytesIO(raw_data))
                text = "".join(page.extract_text() for page in pdf_reader.pages)
                return pd.DataFrame({"content": text.splitlines()})

            # Compressed formats
            elif ext == 'zip':
                with zipfile.ZipFile(BytesIO(raw_data)) as z:
                    first_file = z.namelist()[0]
                    return load_data(BytesIO(z.read(first_file)))
            elif ext == 'tar' or ext == 'tar.gz':
                with tarfile.open(fileobj=BytesIO(raw_data)) as tar:
                    first_file = tar.getmembers()[0]
                    return load_data(BytesIO(tar.extractfile(first_file).read()))
            elif ext == 'gz':
                with gzip.GzipFile(fileobj=BytesIO(raw_data)) as gz:
                    return pd.DataFrame({"content": gz.read().decode('utf-8').splitlines()})
            elif ext == 'bz2':
                with bz2.BZ2File(BytesIO(raw_data)) as bz:
                    return pd.DataFrame({"content": bz.read().decode('utf-8').splitlines()})
            elif ext == 'xz':
                with lzma.LZMAFile(BytesIO(raw_data)) as xz:
                    return pd.DataFrame({"content": xz.read().decode('utf-8').splitlines()})

            # Fallback for unknown formats
            else:
                if fallback_mode == "Text":
                    try:
                        encoding = chardet.detect(raw_data)['encoding'] or 'utf-8'
                        text = raw_data.decode(encoding)
                        return pd.DataFrame({"content": text.splitlines()})
                    except:
                        return pd.DataFrame({"binary": [base64.b64encode(raw_data).decode('utf-8')]})

                elif fallback_mode == "Binary":
                    return pd.DataFrame({"binary": [base64.b64encode(raw_data).decode('utf-8')]})
                else:  # Skip
                    st.warning(f"Skipping unsupported file: {file.name}")
                    return None
        except Exception as e:
            st.error(f"Error loading {file.name}: {str(e)}")
            return pd.DataFrame({"error": [str(e)]})

# Process uploaded files
if files:
    if st.session_state.raw_df is None:
        with st.spinner("Assembling data <span class='spinner'>🔨</span>"):
            valid_files = [f for f in files if f is not None]
            dfs = [load_data(f) for f in valid_files if load_data(f) is not None]
            if dfs:
                st.session_state.raw_df = pd.concat(dfs, ignore_index=True) if len(dfs) > 1 else dfs[0]

    # Input Data Panel
    st.markdown('<div class="panel">', unsafe_allow_html=True)
    st.subheader("📥 Raw Data")
    st.dataframe(st.session_state.raw_df.head(), use_container_width=True)
    with st.expander("File Metadata"):
        st.json(st.session_state.file_metadata)
    st.markdown('</div>', unsafe_allow_html=True)

    # Processing Panel
    st.markdown('<div class="panel action-panel">', unsafe_allow_html=True)
    st.subheader("🔧 Data Forge")
    
    with st.expander("⚙ Processing Controls", expanded=True):
        col1, col2 = st.columns(2)
        with col1:
            cols_to_keep = st.multiselect(
                "Columns to Keep",
                options=st.session_state.raw_df.columns,
                default=list(st.session_state.raw_df.columns)
            )
            rename_cols = {col: st.text_input(f"Rename '{col}'", col, key=f"rn_{col}") 
                         for col in cols_to_keep}
            group_by = st.multiselect("Group By", cols_to_keep)
            agg_func = st.selectbox("Aggregation", ["mean", "sum", "count", "min", "max", "median"]) if group_by else None

            na_strategy = st.selectbox("Handle Missing", ["Keep", "Drop rows", "Drop cols", "Fill", "Interpolate"])
            if na_strategy == "Fill":
                fill_method = st.selectbox("Fill Method", ["Value", "Mean", "Median", "Mode", "FFill", "BFill"])
                fill_value = st.text_input("Fill Value", "0") if fill_method == "Value" else None
        
        with col2:
            outlier_strategy = st.selectbox("Outlier Treatment", ["None", "IQR", "Z-score", "Custom"])
            if outlier_strategy == "Custom":
                outlier_expr = st.text_input("Outlier Condition")
            
            filter_query = st.text_input("Filter (Pandas Query)")
            sort_cols = st.multiselect("Sort Columns", cols_to_keep)
            sort_orders = {col: st.radio(f"{col} Order", ["Asc", "Desc"], key=f"so_{col}") 
                         for col in sort_cols}
            
            calc_cols = st.number_input("New Columns", 0, 10, 0)
            calc_definitions = {}
            for i in range(calc_cols):
                col_name = st.text_input(f"New Col {i+1} Name", key=f"cn_{i}")
                col_expr = st.text_input(f"New Col {i+1} Expr", key=f"ce_{i}")
                if col_name and col_expr:
                    calc_definitions[col_name] = col_expr

            # Text Processing
            text_col = st.selectbox("Text Column", [""] + list(st.session_state.raw_df.columns))
            if text_col:
                text_ops = st.multiselect("Text Ops", ["Lower", "Upper", "Strip", "Split", "Replace"])
                if "Split" in text_ops:
                    split_delim = st.text_input("Split Delimiter", ",")
                if "Replace" in text_ops:
                    replace_from = st.text_input("Replace From")
                    replace_to = st.text_input("Replace To")

    # Action Buttons
    col1, col2, col3, col4, col5 = st.columns(5)
    with col1:
        if st.button("▶ Forge"):
            with st.spinner("Forging <span class='spinner'>⚡</span>"):
                time.sleep(0.7)
                df = st.session_state.raw_df.copy()
                changes = []
                st.session_state.undo_stack.append(df.copy())

                df = df[cols_to_keep].rename(columns={k: v for k, v in rename_cols.items() if k != v})
                if group_by:
                    df = df.groupby(group_by).agg(agg_func).reset_index()
                    changes.append(f"Grouped by {group_by} with {agg_func}")
                
                if na_strategy != "Keep":
                    if na_strategy == "Drop rows":
                        df = df.dropna()
                    elif na_strategy == "Drop cols":
                        df = df.dropna(axis=1)
                    elif na_strategy == "Fill":
                        if fill_method == "Value":
                            df = df.fillna(fill_value)
                        elif fill_method in ["Mean", "Median", "Mode"]:
                            for col in df.select_dtypes(np.number):
                                df[col] = df[col].fillna(getattr(df[col], fill_method.lower())())
                        else:
                            df = df.fillna(method=fill_method.lower())
                    elif na_strategy == "Interpolate":
                        df = df.interpolate()
                    changes.append(f"Handled missing with {na_strategy}")
                
                if outlier_strategy != "None":
                    if outlier_strategy in ["IQR", "Z-score"]:
                        for col in df.select_dtypes(np.number):
                            if outlier_strategy == "IQR":
                                Q1, Q3 = df[col].quantile([0.25, 0.75])
                                IQR = Q3 - Q1
                                mask = (df[col] >= Q1 - 1.5*IQR) & (df[col] <= Q3 + 1.5*IQR)
                            else:
                                z = np.abs((df[col] - df[col].mean()) / df[col].std())
                                mask = z < 3
                            df = df[mask]
                    elif outlier_strategy == "Custom" and outlier_expr:
                        df = df.query(outlier_expr)
                    changes.append(f"Outliers treated with {outlier_strategy}")
                
                if filter_query:
                    df = df.query(filter_query)
                    changes.append(f"Filtered: {filter_query}")
                if sort_cols:
                    df = df.sort_values(sort_cols, ascending=[o == "Asc" for o in sort_orders.values()])
                    changes.append(f"Sorted by {sort_cols}")
                for name, expr in calc_definitions.items():
                    df[name] = pd.eval(expr, target=df)
                    changes.append(f"Added column {name}")
                
                if text_col and text_ops:
                    if "Lower" in text_ops:
                        df[text_col] = df[text_col].str.lower()
                    if "Upper" in text_ops:
                        df[text_col] = df[text_col].str.upper()
                    if "Strip" in text_ops:
                        df[text_col] = df[text_col].str.strip()
                    if "Split" in text_ops:
                        df[text_col] = df[text_col].str.split(split_delim)
                    if "Replace" in text_ops and replace_from and replace_to:
                        df[text_col] = df[text_col].str.replace(replace_from, replace_to)
                    changes.append(f"Text ops on {text_col}: {text_ops}")

                st.session_state.processed_df = df
                st.session_state.history.append({
                    "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                    "changes": changes
                })
                st.success("Data forged successfully!", icon="🔥")

    with col2:
        if st.button("↺ Undo", disabled=not st.session_state.undo_stack):
            with st.spinner("Undoing <span class='spinner'>⟲</span>"):
                time.sleep(0.5)
                st.session_state.processed_df = st.session_state.undo_stack.pop()
                st.session_state.history.append({
                    "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                    "changes": ["Undo"]
                })
                st.success("Undo completed!", icon="↩️")

    with col3:
        preset_name = st.text_input("Preset Name", key="preset_input")
        if st.button("💾 Save Preset") and preset_name:
            st.session_state.presets[preset_name] = {
                "cols": cols_to_keep,
                "rename": rename_cols,
                "group_by": group_by,
                "agg_func": agg_func,
                "na_strategy": na_strategy,
                "text_ops": text_ops if text_col else []
            }
            st.success(f"Preset '{preset_name}' saved!", icon="💾")

    with col4:
        preset_load = st.selectbox("Load Preset", [""] + list(st.session_state.presets.keys()))
        if preset_load and st.button("Load Preset"):
            preset = st.session_state.presets[preset_load]
            st.session_state.raw_df = st.session_state.raw_df[preset["cols"]].rename(
                columns={k: v for k, v in preset["rename"].items() if k != v}
            )
            st.success(f"Preset '{preset_load}' loaded!", icon="📂")

    with col5:
        if st.button("🗑 Clear All"):
            with st.spinner("Clearing <span class='spinner'>🗑️</span>"):
                time.sleep(0.5)
                st.session_state.raw_df = None
                st.session_state.processed_df = None
                st.session_state.undo_stack = []
                st.session_state.history = []
                st.session_state.file_metadata = {}
                st.success("All cleared!", icon="🧹")

    if st.session_state.processed_df is not None:
        st.subheader("🔍 Forged Data")
        st.dataframe(st.session_state.processed_df.head(), use_container_width=True)
    
    st.markdown('</div>', unsafe_allow_html=True)

    # Visualization Panel
    st.markdown('<div class="panel">', unsafe_allow_html=True)
    st.subheader("📊 Visualization Forge")
    if st.session_state.processed_df is not None:
        viz_type = st.selectbox("Chart Type", 
                              ["Scatter", "Line", "Bar", "Histogram", "Box", "Pie", "Heatmap", "Pairplot", "Violin", "Area"])
        
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            x_col = st.selectbox("X-Axis", st.session_state.processed_df.columns)
        with col2:
            y_col = st.selectbox("Y-Axis", 
                               st.session_state.processed_df.columns if viz_type not in ["Histogram", "Pie"] else [""])
        with col3:
            color_col = st.selectbox("Color By", [""] + list(st.session_state.processed_df.columns))
        with col4:
            size_col = st.selectbox("Size By", [""] + list(st.session_state.processed_df.select_dtypes(np.number).columns))
        
        with st.spinner("Visualizing <span class='spinner'>🎨</span>"):
            time.sleep(0.5)
            fig_kwargs = {"data_frame": st.session_state.processed_df, "x": x_col}
            if y_col:
                fig_kwargs["y"] = y_col
            if color_col:
                fig_kwargs["color"] = color_col
            if size_col:
                fig_kwargs["size"] = size_col

            if viz_type == "Scatter":
                fig = px.scatter(**fig_kwargs)
            elif viz_type == "Line":
                fig = px.line(**fig_kwargs)
            elif viz_type == "Bar":
                fig = px.bar(**fig_kwargs)
            elif viz_type == "Histogram":
                fig = px.histogram(**fig_kwargs)
            elif viz_type == "Box":
                fig = px.box(**fig_kwargs)
            elif viz_type == "Pie":
                fig = px.pie(names=x_col, values=y_col if y_col else None, **fig_kwargs)
            elif viz_type == "Heatmap":
                fig = px.density_heatmap(**fig_kwargs)
            elif viz_type == "Pairplot":
                pair_fig = sns.pairplot(st.session_state.processed_df.select_dtypes(np.number))
                st.pyplot(pair_fig.figure)
            elif viz_type == "Violin":
                fig = px.violin(**fig_kwargs)
            elif viz_type == "Area":
                fig = px.area(**fig_kwargs)
            
            if viz_type != "Pairplot":
                fig.update_layout(transition_duration=500, height=600, plot_bgcolor="#1f2a44", 
                                paper_bgcolor="#1f2a44", font_color="#e2e8f0")
                st.plotly_chart(fig, use_container_width=True)
    
    st.markdown('</div>', unsafe_allow_html=True)

    # Export & Analysis Panel
    st.markdown('<div class="panel">', unsafe_allow_html=True)
    st.subheader("💾 Export & Insights")
    if st.session_state.processed_df is not None:
        col1, col2 = st.columns(2)
        with col1:
            export_formats = [
                "CSV", "Excel", "JSON", "Parquet", "SQL", "Word", "PDF", "TXT", "XML", "HTML",
                "Markdown", "Feather", "Pickle", "HDF5", "YAML", "TOML", "SQLite", "ZIP", "GZIP", "BZ2", "XZ"
            ]
            export_format = st.selectbox("Export As", export_formats)
            export_name = st.text_input("File Name", "dataforge_omni")
            compression = st.checkbox("Compress Output", value=False) if export_format not in ["ZIP", "GZIP", "BZ2", "XZ"] else True
            
            if st.button("📤 Export"):
                with st.spinner("Exporting <span class='spinner'>↓</span>"):
                    time.sleep(0.7)
                    output = BytesIO()
                    
                    if export_format == "CSV":
                        data = st.session_state.processed_df.to_csv(index=False)
                        ext = "csv"
                        mime = "text/csv"
                    elif export_format == "Excel":
                        with pd.ExcelWriter(output) as writer:
                            st.session_state.processed_df.to_excel(writer, index=False)
                        data = output.getvalue()
                        ext = "xlsx"
                        mime = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                    elif export_format == "JSON":
                        data = st.session_state.processed_df.to_json()
                        ext = "json"
                        mime = "application/json"
                    elif export_format == "Parquet":
                        data = st.session_state.processed_df.to_parquet()
                        ext = "parquet"
                        mime = "application/octet-stream"
                    elif export_format == "SQL":
                        conn = sqlite3.connect(':memory:')
                        st.session_state.processed_df.to_sql("data", conn, index=False)
                        data = conn.cursor().execute("SELECT sql FROM sqlite_master WHERE type='table'").fetchall()[0][0]
                        ext = "sql"
                        mime = "text/plain"
                    elif export_format == "Word":
                        doc = Document()
                        table = doc.add_table(rows=1, cols=len(st.session_state.processed_df.columns))
                        hdr_cells = table.rows[0].cells
                        for i, col in enumerate(st.session_state.processed_df.columns):
                            hdr_cells[i].text = col
                        for _, row in st.session_state.processed_df.iterrows():
                            row_cells = table.add_row().cells
                            for i, val in enumerate(row):
                                row_cells[i].text = str(val)
                        doc.save(output)
                        data = output.getvalue()
                        ext = "docx"
                        mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    elif export_format == "PDF":
                        pdf = SimpleDocTemplate(output, pagesize=letter)
                        data_list = [st.session_state.processed_df.columns.tolist()] + st.session_state.processed_df.values.tolist()
                        table = Table(data_list)
                        pdf.build([table])
                        data = output.getvalue()
                        ext = "pdf"
                        mime = "application/pdf"
                    elif export_format == "TXT":
                        data = st.session_state.processed_df.to_string(index=False)
                        ext = "txt"
                        mime = "text/plain"
                    elif export_format == "XML":
                        data = st.session_state.processed_df.to_xml()
                        ext = "xml"
                        mime = "application/xml"
                    elif export_format == "HTML":
                        data = st.session_state.processed_df.to_html()
                        ext = "html"
                        mime = "text/html"
                    elif export_format == "Markdown":
                        data = st.session_state.processed_df.to_markdown()
                        ext = "md"
                        mime = "text/markdown"
                    elif export_format == "Feather":
                        feather.write_dataframe(st.session_state.processed_df, output)
                        data = output.getvalue()
                        ext = "feather"
                        mime = "application/octet-stream"
                    elif export_format == "Pickle":
                        data = pickle.dumps(st.session_state.processed_df)
                        ext = "pkl"
                        mime = "application/octet-stream"
                    elif export_format == "HDF5":
                        st.session_state.processed_df.to_hdf(output, key='data', mode='w')
                        data = output.getvalue()
                        ext = "h5"
                        mime = "application/octet-stream"
                    elif export_format == "YAML":
                        data = yaml.dump(st.session_state.processed_df.to_dict(orient='records'))
                        ext = "yaml"
                        mime = "text/yaml"
                    elif export_format == "TOML":
                        data = toml.dumps(st.session_state.processed_df.to_dict(orient='records'))
                        ext = "toml"
                        mime = "text/toml"
                    elif export_format == "SQLite":
                        conn = sqlite3.connect(output)
                        st.session_state.processed_df.to_sql("data", conn, index=False)
                        conn.commit()
                        data = output.getvalue()
                        ext = "db"
                        mime = "application/x-sqlite3"
                    elif export_format == "ZIP":
                        with zipfile.ZipFile(output, 'w', zipfile.ZIP_DEFLATED) as zf:
                            zf.writestr(f"{export_name}.csv", st.session_state.processed_df.to_csv(index=False))
                        data = output.getvalue()
                        ext = "zip"
                        mime = "application/zip"
                    elif export_format == "GZIP":
                        with gzip.GzipFile(fileobj=output, mode='wb') as gz:
                            gz.write(st.session_state.processed_df.to_csv(index=False).encode('utf-8'))
                        data = output.getvalue()
                        ext = "gz"
                        mime = "application/gzip"
                    elif export_format == "BZ2":
                        with bz2.BZ2File(output, 'wb') as bz:
                            bz.write(st.session_state.processed_df.to_csv(index=False).encode('utf-8'))
                        data = output.getvalue()
                        ext = "bz2"
                        mime = "application/x-bzip2"
                    elif export_format == "XZ":
                        with lzma.LZMAFile(output, 'wb') as xz:
                            xz.write(st.session_state.processed_df.to_csv(index=False).encode('utf-8'))
                        data = output.getvalue()
                        ext = "xz"
                        mime = "application/x-xz"

                    if compression and export_format not in ["ZIP", "GZIP", "BZ2", "XZ"]:
                        with gzip.GzipFile(fileobj=output, mode='wb') as gz:
                            gz.write(data.encode('utf-8') if isinstance(data, str) else data)
                        data = output.getvalue()
                        ext = f"{ext}.gz"
                        mime = "application/gzip"

                    bin_str = base64.b64encode(data.encode('utf-8') if isinstance(data, str) else data).decode()
                    st.markdown(
                        f'<a href="data:{mime};base64,{bin_str}" download="{export_name}.{ext}" style="color: #a78bfa;" class="glow-text">📥 Download {export_name}.{ext}</a>',
                        unsafe_allow_html=True
                    )

        with col2:
            with st.expander("🔍 Insights & History"):
                if st.button("Generate Insights"):
                    with st.spinner("Analyzing <span class='spinner'>🔍</span>"):
                        time.sleep(0.5)
                        insights = []
                        for col in st.session_state.processed_df.columns:
                            if st.session_state.processed_df[col].dtype in [np.float64, np.int64]:
                                skew = st.session_state.processed_df[col].skew()
                                kurt = st.session_state.processed_df[col].kurtosis()
                                if abs(skew) > 1:
                                    insights.append(f"{col}: Skewed ({skew:.2f})")
                                if abs(kurt) > 3:
                                    insights.append(f"{col}: High kurtosis ({kurt:.2f})")
                            missing = st.session_state.processed_df[col].isna().mean()
                            if missing > 0:
                                insights.append(f"{col}: {missing*100:.1f}% missing")
                        st.write(insights if insights else ["No significant issues detected"])
                
                st.subheader("Recent History")
                for entry in st.session_state.history[-5:]:
                    st.write(f"{entry['timestamp']}: {', '.join(entry['changes'])}")
    
    st.markdown('</div>', unsafe_allow_html=True)

else:
    st.markdown('<div class="panel">', unsafe_allow_html=True)
    st.info("Upload any file to begin forging your data universe!")
    st.markdown('</div>', unsafe_allow_html=True)

# Footer
st.markdown("""
    <div style='text-align: center; color: #a78bfa; padding: 30px;'>
        <p class="glow-text"><strong>DataForge Omni:</strong> Every format, every possibility</p>
        <p>Universal file support | Advanced transformations | Dynamic visualizations</p>
    </div>
""", unsafe_allow_html=True)
st.markdown('</div>', unsafe_allow_html=True)
